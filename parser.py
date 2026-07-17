import os
from pypdf import PdfReader
from docx import Document
from PIL import Image
import pytesseract
import pypdfium2 as pdfium

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def _ocr_pdf_fallback(file_path: str) -> str:
    """Taranmış PDF'leri OCR ile metne çevirir."""
    pdf = None
    try:
        pdf = pdfium.PdfDocument(file_path)
        text = ""
        for i in range(len(pdf)):
            page = pdf[i]
            bitmap = page.render(scale=2)
            pil_image = bitmap.to_pil()
            page_text = pytesseract.image_to_string(pil_image, lang="tur+eng")
            text += page_text + "\n"
        return text.strip()
    except Exception as e:
        return f"PDF OCR hatası: {str(e)}"
    finally:
        if pdf:
            pdf.close()


def extract_text_from_file(file_path: str) -> str:
    if not os.path.exists(file_path):
        return "Hata: Dosya bulunamadı!"

    file_extension = file_path.lower().split('.')[-1]
    text = ""

    if file_extension == 'pdf':
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        except Exception as e:
            return f"PDF okuma hatası: {str(e)}"

        if len(text.strip()) < 50:
            text = _ocr_pdf_fallback(file_path)

    elif file_extension == 'docx':
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            return f"DOCX okuma hatası: {str(e)}"

    else:
        return "Hata: Desteklenmeyen dosya formatı! Lütfen PDF veya DOCX yükleyin."

    return text.strip()

# --- LOKAL TEST İÇİN ---
if __name__ == "__main__":
    ornek_dosya = "ornek_cv.pdf" 
    if os.path.exists(ornek_dosya):
        cikti = extract_text_from_file(ornek_dosya)
        print("--- ÇIKARILAN METİN ÖZETİ ---")
        print(cikti[:300] + "...")
    else:
        print("ℹ️ Test için klasörde 'ornek_cv.pdf' bulunamadı, ancak parser fonksiyona hazır.")